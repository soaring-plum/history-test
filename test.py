import tkinter
import docx
import tkinter.messagebox
import sys

def tianxitx():
    tkinter.messagebox.showinfo('made by 湉曦tx', '感谢你的使用！\n由园艺202 李翔倾情奉上。\nqq:1914076211(男生就算了)')


def ex():
    exit()


def judge():
    z1 = ''
    if status1.get() == 1:
        z1 = z1 + 'A'
    if status2.get() == 1:
        z1 = z1 + 'B'
    if status3.get() == 1:
        z1 = z1 + 'C'
    if status4.get() == 1:
        z1 = z1 + 'D'
    if z1 == a[n]:
        s2.set('选了' + z1 + '对')
    else:
        s2.set('选了' + z1 + '错')
    checkbox1.deselect()
    checkbox2.deselect()
    checkbox3.deselect()
    checkbox4.deselect()


# def jump():
#     pass


def tip():
    tkinter.messagebox.showinfo('使用说明', '判断题A为T，B为F。')


def b1():
    button1 = tkinter.Button(root, text='next', bg='Pink', command=nt)
    button1.place(x=100, y=300, width=200, height=100)


def b2():
    button2 = tkinter.Button(root, text='preview', bg='Cyan', command=pr)
    button2.place(x=300, y=300, width=200, height=100)


def b3():
    button3 = tkinter.Button(root, text='提交', bg='Red', command=judge)
    button3.place(x=50, y=300, width=50, height=100)


def b4():
    button4 = tkinter.Button(root, text='made by 湉曦tx', bg='HotPink', command=tianxitx)
    button4.place(x=700, y=300, width=200, height=100)


def b5():
    button5 = tkinter.Button(root, text='使用说明', bg='Violet', command=tip)
    button5.place(x=500, y=300, width=200, height=100)


def b6():
    button6 = tkinter.Button(root, text='退出', bg='Orchid', command=sys.exit)
    button6.place(x=1000, y=300, width=200, height=100)


# def b7():
#     button7 = tkinter.Button(root, text='跳转', bg='Pink', command=jump)
#     button7.pack()


def l1():
    global s1
    label1 = tkinter.Label(root, textvariable=s1, justify='left', anchor='nw', font=('', 15, ''), wraplength='1200',
                           height=10)
    label1.place(x=0, y=50)


def l2():
    global s2
    label2 = tkinter.Label(root, textvariable=s2, justify='left', anchor='nw', width=1920, font=('', 15, ''), fg="Red")
    label2.pack()


# def t1():
#     text1 = tkinter.Text(root)
#     text1.pack()


def pr():
    global n
    n = n - 1
    s1.set(q[n])
    s2.set('')
    checkbox1.deselect()
    checkbox2.deselect()
    checkbox3.deselect()
    checkbox4.deselect()


def nt():
    global n
    n = n + 1
    s1.set(q[n])
    s2.set('')
    checkbox1.deselect()
    checkbox2.deselect()
    checkbox3.deselect()
    checkbox4.deselect()


def qu1(list1, file1, a1):
    d = 1
    t = ''
    for i in file1.paragraphs:
        if '多项选择题' in i.text:
            list1.append(t)
            break
        if str(d) + '、' in i.text:
            d = d + 1
            list1.append(t)
            if 'A' in i.text:
                t = i.text
                t = t.replace('A', '')
                a1.append('A')
            elif 'B' in i.text:
                t = i.text
                t = t.replace('B', '')
                a1.append('B')
            elif 'C' in i.text:
                t = i.text
                t = t.replace('C', '')
                a1.append('C')
            elif 'D' in i.text:
                t = i.text
                t = t.replace('D', '')
                a1.append('D')
            t = t + '\n'
        else:
            z = i.text
            z = "".join(z.split())
            z = z.replace('B', '\nB')
            z = z.replace('C', '\nC')
            z = z.replace('D', '\nD')
            t = t + z


def qu2(list2, file2, a2):
    d = 1
    t = ''
    t0 = False
    for i in file2.paragraphs:
        if '多项选择题' in i.text:
            t0 = True
        if '判断题' in i.text:
            list2.append(t)
            break
        if t0:
            if str(d) + '、' in i.text:
                d = d + 1
                list2.append(t)
                a4 = ''
                t = i.text
                if 'A' in i.text:
                    t = t.replace('A', '')
                    a4 = a4 + 'A'
                if 'B' in i.text:
                    t = t.replace('B', '')
                    a4 = a4 + 'B'
                if 'C' in i.text:
                    t = t.replace('C', '')
                    a4 = a4 + 'C'
                if 'D' in i.text:
                    t = t.replace('D', '')
                    a4 = a4 + 'D'
                a2.append(a4)
                t = t + '\n'
            else:
                z = i.text
                z = "".join(z.split())
                z = z.replace('B', '\nB')
                z = z.replace('C', '\nC')
                z = z.replace('D', '\nD')
                t = t + z


def qu3(list3, file3, a3):
    t0 = False
    for i in file3.paragraphs:
        if '判断题' in i.text:
            t0 = True
        if t0:
            if 'T' in i.text:
                t = i.text
                t = t.replace('T', '')
                list3.append(t)
                a3.append('A')
            if 'F' in i.text:
                t = i.text
                t = t.replace('F', '')
                list3.append(t)
                a3.append('B')


file = docx.Document('12.docx')
root = tkinter.Tk()
root.state("zoomed")
s1 = tkinter.StringVar()
s1.set('''1、（ ）是浙江文化成熟的标志，它使浙江文化成为全国的主流文化，影响和引领当时中国的潮流。
A、古越文化
B、吴越文化
C、南宋浙学
D、近代海派文化
''')
s2 = tkinter.StringVar()
s2.set('')
status1 = tkinter.IntVar()
status2 = tkinter.IntVar()
status3 = tkinter.IntVar()
status4 = tkinter.IntVar()
n = 0
d1 = 1
d2 = 1
c1 = False
c2 = False
q1 = []
q2 = []
q3 = []
z = []
a = []
qu1(q1, file, a)
qu2(q2, file, a)
qu3(q3, file, a)
q = q1[1:] + q2[1:] + q3
l1()
l2()
checkbox1 = tkinter.Checkbutton(root, text='A', variable=status1)
checkbox1.place(x=5, y=300)
checkbox2 = tkinter.Checkbutton(root, text='B', variable=status2)
checkbox2.place(x=5, y=325)
checkbox3 = tkinter.Checkbutton(root, text='C', variable=status3)
checkbox3.place(x=5, y=350)
checkbox4 = tkinter.Checkbutton(root, text='D', variable=status4)
checkbox4.place(x=5, y=375)
# t1()
b1()
b2()
b3()
b4()
b5()
# b7()
b6()
print(q1)
print(q2)
print(q3)
print(q)
print(a)
# tkinter.messagebox.showinfo('本章含有图片题，请注意！','本章含有图片题，请注意！')
root.update()
root.mainloop()
